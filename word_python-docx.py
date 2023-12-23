import docx

document =docx.Document("C:\\Users\\x\\Desktop\\demo.docx")

print(document.paragraphs) # list of paragraph 

print(document.paragraphs[1].text)

runs = document.paragraphs[1]
print(runs.runs) # each paragrapgh is divided into runs

print(runs.runs[0].text) # run changes when style changes

print(runs.runs[1].bold) #checks for bold and returns a boolean

runs.runs[3].underline = True
runs.runs[3] = "italic and underline"

print(document.paragraphs[1].style)

document.paragraphs[2].style = 'Title'

document.save("C:\\Users\\x\\Desktop\\demo1.docx")

# creating a new word document
d = docx.Document()
d.add_paragraph('This is a paragraph')
d.add_paragraph('This is another paragraph')
d.save("C:\\Users\\x\\Desktop\\demo2.docx")
p = d.paragraphs[0]
p.add_run("This si a new run")
d.save("C:\\Users\\x\\Desktop\\demo2.docx")
print(p.runs)
p.runs[1].bold = True
d.save("C:\\Users\\x\\Desktop\\demo2.docx")




